# app/document_filler.py
import fillpdf
import fitz
from docx import Document
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import os
from typing import Dict, List
import tempfile

class DocumentFiller:
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    def fill_pdf_form(self, input_path: str, output_path: str, 
                     field_data: Dict[str, str]) -> bool:
        """Fill PDF form using fillpdf library"""
        try:
            # Method 1: Using fillpdf (recommended for form fields)
            fillpdf.write_fillable_pdf(input_path, output_path, field_data)
            return True
            
        except Exception as e:
            print(f"fillpdf failed: {e}, trying PyMuPDF method")
            return self._fill_pdf_pymupdf(input_path, output_path, field_data)
    
    def _fill_pdf_pymupdf(self, input_path: str, output_path: str, 
                         field_data: Dict[str, str]) -> bool:
        """Fill PDF using PyMuPDF as fallback"""
        try:
            doc = fitz.open(input_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                widgets = page.widgets()
                
                for widget in widgets:
                    field_name = widget.field_name
                    if field_name in field_data:
                        widget.field_value = field_data[field_name]
                        widget.update()
            
            doc.save(output_path)
            doc.close()
            return True
            
        except Exception as e:
            print(f"PyMuPDF fill failed: {e}")
            return self._fill_pdf_overlay(input_path, output_path, field_data)
    
    def _fill_pdf_overlay(self, input_path: str, output_path: str, 
                         field_data: Dict[str, str]) -> bool:
        """Create overlay PDF with text (last resort)"""
        try:
            # This is a simplified overlay method
            # In production, you'd need more sophisticated positioning
            doc = fitz.open(input_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Add text overlays (you'd need to determine positions)
                y_position = 750  # Start from top
                for field_name, value in field_data.items():
                    if value and len(value) > 0:
                        # Insert text at calculated position
                        text_rect = fitz.Rect(50, y_position, 500, y_position + 20)
                        page.insert_text(text_rect.tl, f"{field_name}: {value}", 
                                       fontsize=10, color=(0, 0, 0))
                        y_position -= 25
            
            doc.save(output_path)
            doc.close()
            return True
            
        except Exception as e:
            print(f"PDF overlay failed: {e}")
            return False
    
    def fill_word_document(self, template_path: str, output_path: str, 
                          replacements: Dict[str, str]) -> bool:
        """Fill Word document with replacements"""
        try:
            doc = Document(template_path)
            
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, replacement in replacements.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, replacement)
            
            # Replace in headers and footers
            for section in doc.sections:
                header = section.header
                footer = section.footer
                
                for paragraph in header.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
                
                for paragraph in footer.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error filling Word document: {e}")
            return False
    
    def fill_excel_template(self, template_path: str, output_path: str, 
                           cell_data: Dict[str, str]) -> bool:
        """Fill Excel template with data"""
        try:
            workbook = openpyxl.load_workbook(template_path)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Fill specific cells
                for cell_ref, value in cell_data.items():
                    try:
                        sheet[cell_ref] = value
                    except Exception as e:
                        print(f"Error filling cell {cell_ref}: {e}")
                
                # Search and replace in all cells
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for placeholder, replacement in cell_data.items():
                                if placeholder in str(cell.value):
                                    cell.value = str(cell.value).replace(placeholder, replacement)
            
            workbook.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error filling Excel template: {e}")
            return False
    
    def create_response_document(self, responses: Dict[str, str], 
                               output_path: str, format_type: str = 'docx') -> bool:
        """Create a new document with all responses"""
        try:
            if format_type.lower() == 'docx':
                return self._create_word_response(responses, output_path)
            elif format_type.lower() == 'pdf':
                return self._create_pdf_response(responses, output_path)
            else:
                return False
                
        except Exception as e:
            print(f"Error creating response document: {e}")
            return False
    
    def _create_word_response(self, responses: Dict[str, str], output_path: str) -> bool:
        """Create Word document with responses"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Tender Response Document', 0)
            
            # Add responses
            for i, (question, answer) in enumerate(responses.items(), 1):
                # Add question as heading
                doc.add_heading(f'{i}. {question}', level=2)
                
                # Add answer as paragraph
                doc.add_paragraph(answer)
                
                # Add spacing
                doc.add_paragraph()
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating Word response: {e}")
            return False
    
    def _create_pdf_response(self, responses: Dict[str, str], output_path: str) -> bool:
        """Create PDF document with responses"""
        try:
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            y_position = height - 50
            
            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y_position, "Tender Response Document")
            y_position -= 40
            
            # Responses
            for i, (question, answer) in enumerate(responses.items(), 1):
                # Check if we need a new page
                if y_position < 100:
                    c.showPage()
                    y_position = height - 50
                
                # Question
                c.setFont("Helvetica-Bold", 12)
                c.drawString(50, y_position, f"{i}. {question[:80]}...")
                y_position -= 20
                
                # Answer (word wrap)
                c.setFont("Helvetica", 10)
                words = answer.split()
                line = ""
                for word in words:
                    if len(line + word) < 80:
                        line += word + " "
                    else:
                        c.drawString(50, y_position, line)
                        y_position -= 15
                        line = word + " "
                        
                        if y_position < 100:
                            c.showPage()
                            y_position = height - 50
                
                if line:
                    c.drawString(50, y_position, line)
                    y_position -= 30
            
            c.save()
            return True
            
        except Exception as e:
            print(f"Error creating PDF response: {e}")
            return False
