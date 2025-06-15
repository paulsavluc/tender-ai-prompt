# app/document_processor.py
import pypdf
import docx
import openpyxl
import fitz
import re
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
import json

@dataclass
class DocumentField:
    name: str
    field_type: str
    value: str = ""
    position: tuple = (0, 0)
    required: bool = False

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.txt']
    
    def extract_pdf_fields(self, pdf_path: str) -> List[DocumentField]:
        """Extract form fields from PDF using pypdf"""
        fields = []
        try:
            with open(pdf_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                
                # Check if PDF has form fields
                if reader.trailer.get('/Root', {}).get('/AcroForm'):
                    form = reader.trailer['/Root']['/AcroForm']
                    if '/Fields' in form:
                        for field_ref in form['/Fields']:
                            field = field_ref.get_object()
                            field_name = field.get('/T', 'Unknown')
                            field_type = field.get('/FT', '/Tx')  # Default to text
                            
                            # Clean field name
                            if hasattr(field_name, 'decode'):
                                field_name = field_name.decode('utf-8')
                            
                            fields.append(DocumentField(
                                name=str(field_name),
                                field_type=str(field_type),
                                required='/Ff' in field and field['/Ff'] & 2
                            ))
                
                # If no form fields, extract text and identify potential fields
                if not fields:
                    fields = self._extract_text_fields_from_pdf(pdf_path)
                    
        except Exception as e:
            print(f"Error extracting PDF fields: {e}")
            fields = self._extract_text_fields_from_pdf(pdf_path)
        
        return fields
    
    def _extract_text_fields_from_pdf(self, pdf_path: str) -> List[DocumentField]:
        """Extract potential fields from PDF text content"""
        fields = []
        try:
            doc = fitz.open(pdf_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                full_text += page.get_text()
            
            doc.close()
            
            # Look for common field patterns
            field_patterns = [
                r'([A-Za-z\s]+):\s*_+',  # Label: ______
                r'([A-Za-z\s]+)\s*\[\s*\]',  # Label [ ]
                r'([A-Za-z\s]+):\s*\(\s*\)',  # Label: ( )
                r'Please\s+provide\s+([^.]+)',  # Please provide...
                r'Describe\s+([^.]+)',  # Describe...
                r'What\s+is\s+([^?]+)\?',  # What is...?
                r'How\s+([^?]+)\?',  # How...?
            ]
            
            for pattern in field_patterns:
                matches = re.finditer(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    field_name = match.group(1).strip()
                    if len(field_name) > 3 and len(field_name) < 100:
                        fields.append(DocumentField(
                            name=field_name,
                            field_type='text',
                            required=True
                        ))
        
        except Exception as e:
            print(f"Error extracting text fields: {e}")
        
        return fields
    
    def extract_word_content(self, doc_path: str) -> Dict[str, Any]:
        """Extract content and fields from Word document"""
        try:
            doc = docx.Document(doc_path)
            
            # Extract all text
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract tables
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            # Look for placeholders
            placeholders = self._find_placeholders('\n'.join(full_text))
            
            return {
                'text': '\n'.join(full_text),
                'tables': tables_content,
                'placeholders': placeholders,
                'fields': self._extract_word_fields('\n'.join(full_text))
            }
        
        except Exception as e:
            print(f"Error processing Word document: {e}")
            return {'text': '', 'tables': [], 'placeholders': [], 'fields': []}
    
    def _find_placeholders(self, text: str) -> List[str]:
        """Find placeholder patterns in text"""
        patterns = [
            r'\{\{([^}]+)\}\}',  # {{placeholder}}
            r'\[([^\]]+)\]',     # [placeholder]
            r'<([^>]+)>',        # <placeholder>
            r'_+([A-Za-z\s]+)_+' # ___placeholder___
        ]
        
        placeholders = []
        for pattern in patterns:
            matches = re.findall(pattern, text)
            placeholders.extend(matches)
        
        return list(set(placeholders))
    
    def _extract_word_fields(self, text: str) -> List[DocumentField]:
        """Extract fields from Word document text"""
        fields = []
        
        # Question patterns
        question_patterns = [
            r'(\d+\.?\s*[A-Z][^?]*\?)',  # Numbered questions
            r'([A-Z][^:]*:)\s*$',        # Labels ending with colon
            r'(Please\s+[^.]+\.)',       # Instructions starting with "Please"
            r'(Provide\s+[^.]+\.)',      # Instructions starting with "Provide"
        ]
        
        for pattern in question_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                question = match.group(1).strip()
                if len(question) > 10 and len(question) < 200:
                    fields.append(DocumentField(
                        name=question,
                        field_type='text',
                        required=True
                    ))
        
        return fields
    
    def extract_excel_structure(self, excel_path: str) -> Dict[str, Any]:
        """Extract structure from Excel file"""
        try:
            workbook = openpyxl.load_workbook(excel_path)
            structure = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    sheet_data.append(row)
                
                structure[sheet_name] = {
                    'data': sheet_data,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column
                }
            
            return structure
        
        except Exception as e:
            print(f"Error processing Excel file: {e}")
            return {}
