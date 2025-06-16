# tender_app/utils/document_processor.py
import docx
import PyPDF2
import fitz  # PyMuPDF
from openpyxl import load_workbook
import re
from typing import List, Dict, Any

class DocumentProcessor:
    """Handles extraction of fields from various document formats"""
    
    def __init__(self):
        self.supported_formats = ['docx', 'pdf', 'xlsx']
    
    def extract_fields(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """Extract fillable fields from document based on file type"""
        if file_type == 'docx':
            return self._extract_word_fields(file_path)
        elif file_type == 'pdf':
            return self._extract_pdf_fields(file_path)
        elif file_type == 'xlsx':
            return self._extract_excel_fields(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_word_fields(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract fields from Word document"""
        doc = docx.Document(file_path)
        fields = []
        
        # Extract form fields
        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            # Look for text in brackets or underscores (common field patterns)
            field_patterns = [
                r'\[([^\]]+)\]',  # [Field Name]
                r'_{3,}',         # _____ (underscores)
                r'\{([^}]+)\}',   # {Field Name}
            ]
            
            for pattern in field_patterns:
                matches = re.finditer(pattern, paragraph.text)
                print(f"Matches {matches}")
                for match in matches:
                    fields.append({
                        'field_name': match.group(1) if match.groups() else f"Field_{len(fields)+1}",
                        'field_type': 'text',
                        'position_info': {
                            'paragraph_index': paragraph_index,
                            'start': match.start(),
                            'end': match.end(),
                            'original_text': match.group(0)
                        }
                    })
        
        # Extract table cells that might be fillable
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if not cell.text.strip() or cell.text.strip() in ['', ' ']:
                        fields.append({
                            'field_name': f"Table_{table_idx}_Row_{row_idx}_Cell_{cell_idx}",
                            'field_type': 'table_cell',
                            'position_info': {
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx
                            }
                        })
        
        return fields
    
    def _extract_pdf_fields(self, file_path: str) -> List[Dict[str, Any]]:
      """Extract fields from PDF document"""
      fields = []
      
      # Method 1: Try PyMuPDF for form fields (Correct approach)
      try:
          doc = fitz.open(file_path)
          for page_num in range(len(doc)):
              page = doc.load_page(page_num)
              print(f"Page: {page_num}")
              
              # Correct method: Use page.widgets() iterator
              widget_count = 0
              for widget in page.widgets():
                  widget_count += 1
                  fields.append({
                      'field_name': widget.field_name or f"Field_{len(fields)+1}",
                      'field_type': widget.field_type_string,
                      'field_value': widget.field_value,
                      'position_info': {
                          'page': page_num,
                          'rect': list(widget.rect),
                          'field_type': widget.field_type
                      }
                  })
              
              print(f"Found {widget_count} widgets on page {page_num}")
              
              # Alternative: Check annotations for widget types
              if widget_count == 0:
                  for annot in page.annots():
                      if annot.type[1] == 'Widget':  # Check if annotation is a widget
                          fields.append({
                              'field_name': f"Widget_{len(fields)+1}",
                              'field_type': 'widget_annotation',
                              'position_info': {
                                  'page': page_num,
                                  'rect': list(annot.rect),
                                  'content': annot.content
                              }
                          })
          
          doc.close()
      except Exception as e:
          print(f"Error extracting PDF fields with PyMuPDF: {e}")
      
      # Method 2: PyPDF2 for AcroForm fields
      if not fields:
          try:
              with open(file_path, 'rb') as file:
                  pdf_reader = PyPDF2.PdfReader(file)
                  
                  # Check if PDF has form fields in AcroForm
                  if pdf_reader.trailer.get("/Root", {}).get("/AcroForm"):
                      form_fields = pdf_reader.get_form_text_fields()
                      if form_fields:
                          for field_name, field_value in form_fields.items():
                              fields.append({
                                  'field_name': field_name,
                                  'field_type': 'acroform_text',
                                  'field_value': field_value,
                                  'position_info': {
                                      'source': 'AcroForm'
                                  }
                              })
                      
                      # Also try to get all fields including non-text
                      if hasattr(pdf_reader, 'get_fields'):
                          all_fields = pdf_reader.get_fields()
                          for field_name, field_obj in all_fields.items():
                              if field_name not in [f['field_name'] for f in fields]:
                                  fields.append({
                                      'field_name': field_name,
                                      'field_type': field_obj.get('/FT', 'unknown'),
                                      'field_value': field_obj.get('/V', ''),
                                      'position_info': {
                                          'source': 'AcroForm_All'
                                      }
                                  })
          except Exception as e:
              print(f"Error extracting PDF fields with PyPDF2: {e}")
      
      # Method 3: Text pattern extraction as fallback
      if not fields:
          try:
              with open(file_path, 'rb') as file:
                  pdf_reader = PyPDF2.PdfReader(file)
                  for page_num, page in enumerate(pdf_reader.pages):
                      text = page.extract_text()
                      
                      # Enhanced field patterns
                      field_patterns = [
                          (r'Name:\s*([_\.\s]{3,})', 'name_field'),
                          (r'Date:\s*([_\.\s]{3,})', 'date_field'),
                          (r'Signature:\s*([_\.\s]{3,})', 'signature_field'),
                          (r'Address:\s*([_\.\s]{3,})', 'address_field'),
                          (r'\[([^\]]+)\]', 'bracketed_field'),
                          (r'_{5,}', 'underscore_field'),
                          (r'\.{5,}', 'dotted_field'),
                      ]
                      
                      for pattern, field_type in field_patterns:
                          matches = re.finditer(pattern, text, re.IGNORECASE)
                          for match in matches:
                              field_name = match.group(1) if match.groups() and len(match.groups()) > 0 else f"{field_type}_{len(fields)+1}"
                              fields.append({
                                  'field_name': field_name.strip(),
                                  'field_type': field_type,
                                  'position_info': {
                                      'page': page_num,
                                      'start': match.start(),
                                      'end': match.end(),
                                      'original_text': match.group(0)
                                  }
                              })
          except Exception as e:
              print(f"Error with text pattern extraction: {e}")
      
      return fields

    
    def _extract_excel_fields(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract fields from Excel document"""
        workbook = load_workbook(file_path)
        fields = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            for row in sheet.iter_rows():
                for cell in row:
                    # Look for empty cells or cells with placeholder text
                    if (cell.value is None or 
                        (isinstance(cell.value, str) and 
                         any(placeholder in cell.value.lower() for placeholder in 
                             ['[enter', 'fill in', 'insert', 'add your']))):
                        
                        fields.append({
                            'field_name': f"{sheet_name}_{cell.coordinate}",
                            'field_type': 'cell',
                            'position_info': {
                                'sheet': sheet_name,
                                'coordinate': cell.coordinate,
                                'row': cell.row,
                                'column': cell.column
                            }
                        })
        
        return fields

    def extract_reference_content(self, file_path: str, file_type: str) -> str:
        """Extract text content from reference documents"""
        if file_type == 'docx':
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_type == 'pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ''
                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'
                return text
        elif file_type == 'xlsx':
            workbook = load_workbook(file_path)
            text = ''
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            text += str(cell.value) + ' '
                    text += '\n'
            return text
        else:
            return ''
