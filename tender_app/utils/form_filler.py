# tender_app/utils/form_filler.py
import docx
import fitz  # PyMuPDF
from openpyxl import load_workbook
import re
from typing import Dict, Any
import os

class FormFiller:
    """Fills forms with generated content"""
    
    def fill_document(self, template_path: str, output_path: str, 
                     field_content: Dict[str, str], file_type: str) -> bool:
        """Fill document with generated content"""
        try:
            if file_type == 'docx':
                return self._fill_word_document(template_path, output_path, field_content)
            elif file_type == 'pdf':
                return self._fill_pdf_document(template_path, output_path, field_content)
            elif file_type == 'xlsx':
                return self._fill_excel_document(template_path, output_path, field_content)
            else:
                return False
        except Exception as e:
            print(f"Error filling document: {e}")
            return False
    
    def _fill_word_document(self, template_path: str, output_path: str, 
                          field_content: Dict[str, str]) -> bool:
        """Fill Word document with content"""
        doc = docx.Document(template_path)
        
        # Replace text patterns in paragraphs
        for paragraph in doc.paragraphs:
            for field_name, content in field_content.items():
                # Replace various field patterns
                patterns = [
                    f'[{field_name}]',
                    f'{{{field_name}}}',
                    field_name
                ]
                
                for pattern in patterns:
                    if pattern in paragraph.text:
                        paragraph.text = paragraph.text.replace(pattern, content)
        
        # Fill table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for field_name, content in field_content.items():
                        if field_name in cell.text or not cell.text.strip():
                            cell.text = content
                            break
        
        doc.save(output_path)
        return True
    
    def _fill_pdf_document(self, template_path: str, output_path: str, 
                         field_content: Dict[str, str]) -> bool:
        """Fill PDF document with content"""
        try:
            doc = fitz.open(template_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Fill form fields
                if page.first_widget:
                    widget = page.first_widget
                    while widget:
                        field_name = widget.field_name
                        if field_name in field_content:
                            widget.field_value = field_content[field_name]
                            widget.update()
                        widget = widget.next
            
            doc.save(output_path)
            doc.close()
            return True
            
        except Exception as e:
            print(f"Error filling PDF: {e}")
            return False
    
    def _fill_excel_document(self, template_path: str, output_path: str, 
                           field_content: Dict[str, str]) -> bool:
        """Fill Excel document with content"""
        workbook = load_workbook(template_path)
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            for field_name, content in field_content.items():
                # Parse field name to get sheet and cell info
                if sheet_name in field_name:
                    # Extract coordinate from field name
                    parts = field_name.split('_')
                    if len(parts) >= 2:
                        coordinate = parts[-1]
                        try:
                            sheet[coordinate] = content
                        except:
                            continue
        
        workbook.save(output_path)
        return True
