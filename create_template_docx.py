# create_template_docx.py
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template_docx():
    """Create a template DOCX with various field types"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Employee Information Form Template', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add various field types that your extractor can find
    doc.add_heading('Personal Information', level=1)
    
    # Text placeholders (these will be detected by pattern matching)
    p1 = doc.add_paragraph()
    p1.add_run('Full Name: ').bold = True
    p1.add_run('_' * 30)
    
    p2 = doc.add_paragraph()
    p2.add_run('Date of Birth: ').bold = True
    p2.add_run('_' * 20)
    
    p3 = doc.add_paragraph()
    p3.add_run('Address: ').bold = True
    p3.add_run('_' * 50)
    
    # Dotted fields
    p4 = doc.add_paragraph()
    p4.add_run('Phone Number: ').bold = True
    p4.add_run('.' * 25)
    
    # Bracketed fields
    p5 = doc.add_paragraph()
    p5.add_run('Department: ').bold = True
    p5.add_run('[Select Department]')
    
    p6 = doc.add_paragraph()
    p6.add_run('Position: ').bold = True
    p6.add_run('[Job Title]')
    
    # Employment section
    doc.add_heading('Employment Details', level=1)
    
    p7 = doc.add_paragraph()
    p7.add_run('Start Date: ').bold = True
    p7.add_run('_' * 15)
    
    p8 = doc.add_paragraph()
    p8.add_run('Salary: ').bold = True
    p8.add_run('$').bold = True
    p8.add_run('_' * 15)
    
    p9 = doc.add_paragraph()
    p9.add_run('Manager: ').bold = True
    p9.add_run('[Manager Name]')
    
    # Signature section
    doc.add_heading('Signatures', level=1)
    
    p10 = doc.add_paragraph()
    p10.add_run('Employee Signature: ').bold = True
    p10.add_run('_' * 30)
    
    p11 = doc.add_paragraph()
    p11.add_run('Date: ').bold = True
    p11.add_run('_' * 15)
    
    # Table with fields
    doc.add_heading('Emergency Contacts', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Relationship'
    hdr_cells[2].text = 'Phone'
    
    # Data rows with placeholders
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '[Contact Name 1]'
    row1_cells[1].text = '[Relationship 1]'
    row1_cells[2].text = '_____________'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = '[Contact Name 2]'
    row2_cells[1].text = '[Relationship 2]'
    row2_cells[2].text = '_____________'
    
    doc.save('employee_form_template.docx')
    print("Template DOCX created: employee_form_template.docx")

def create_filled_reference_docx():
    """Create a reference DOCX with filled data"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Employee Information Form - FILLED', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Personal Information
    doc.add_heading('Personal Information', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('Full Name: ').bold = True
    p1.add_run('John Michael Smith')
    
    p2 = doc.add_paragraph()
    p2.add_run('Date of Birth: ').bold = True
    p2.add_run('January 15, 1985')
    
    p3 = doc.add_paragraph()
    p3.add_run('Address: ').bold = True
    p3.add_run('123 Main Street, Anytown, ST 12345')
    
    p4 = doc.add_paragraph()
    p4.add_run('Phone Number: ').bold = True
    p4.add_run('(555) 123-4567')
    
    p5 = doc.add_paragraph()
    p5.add_run('Department: ').bold = True
    p5.add_run('Information Technology')
    
    p6 = doc.add_paragraph()
    p6.add_run('Position: ').bold = True
    p6.add_run('Senior Software Developer')
    
    # Employment Details
    doc.add_heading('Employment Details', level=1)
    
    p7 = doc.add_paragraph()
    p7.add_run('Start Date: ').bold = True
    p7.add_run('March 1, 2023')
    
    p8 = doc.add_paragraph()
    p8.add_run('Salary: ').bold = True
    p8.add_run('$').bold = True
    p8.add_run('75,000 annually')
    
    p9 = doc.add_paragraph()
    p9.add_run('Manager: ').bold = True
    p9.add_run('Sarah Johnson')
    
    # Signatures
    doc.add_heading('Signatures', level=1)
    
    p10 = doc.add_paragraph()
    p10.add_run('Employee Signature: ').bold = True
    p10.add_run('John M. Smith')
    
    p11 = doc.add_paragraph()
    p11.add_run('Date: ').bold = True
    p11.add_run('June 15, 2025')
    
    # Emergency Contacts Table
    doc.add_heading('Emergency Contacts', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Relationship'
    hdr_cells[2].text = 'Phone'
    
    # Filled data rows
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Mary Smith'
    row1_cells[1].text = 'Spouse'
    row1_cells[2].text = '(555) 123-4568'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Robert Smith'
    row2_cells[1].text = 'Father'
    row2_cells[2].text = '(555) 987-6543'
    
    doc.save('employee_form_filled_reference.docx')
    print("Reference DOCX created: employee_form_filled_reference.docx")

def create_advanced_template_with_content_controls():
    """Create template using docx-form approach with content controls"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Advanced Form Template with Content Controls', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Note about content controls
    note = doc.add_paragraph()
    note.add_run('Note: ').bold = True
    note.add_run('This template uses various field patterns that can be detected by the extraction system.')
    
    # Multiple field pattern types for testing
    doc.add_heading('Field Pattern Testing', level=1)
    
    # Pattern 1: Underscore fields
    patterns = [
        ('Name', '_' * 25),
        ('Email', '_' * 30),
        ('Employee ID', '_' * 10),
        ('Hire Date', '_' * 15),
    ]
    
    for label, pattern in patterns:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(pattern)
    
    # Pattern 2: Dotted fields
    doc.add_heading('Dotted Line Fields', level=2)
    dotted_patterns = [
        ('Social Security Number', '.' * 15),
        ('Emergency Contact', '.' * 25),
        ('Preferred Name', '.' * 20),
    ]
    
    for label, pattern in dotted_patterns:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(pattern)
    
    # Pattern 3: Bracketed fields
    doc.add_heading('Selection Fields', level=2)
    bracketed_patterns = [
        ('Gender', '[Male/Female/Other]'),
        ('Marital Status', '[Single/Married/Divorced]'),
        ('Work Location', '[Office/Remote/Hybrid]'),
        ('Benefits Plan', '[Basic/Premium/Executive]'),
    ]
    
    for label, pattern in bracketed_patterns:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(pattern)
    
    # Pattern 4: Mixed patterns
    doc.add_heading('Mixed Pattern Fields', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Signature: ').bold = True
    p.add_run('_' * 30)
    p.add_run('  Date: ').bold = True
    p.add_run('_' * 12)
    
    doc.save('advanced_template.docx')
    print("Advanced template created: advanced_template.docx")

if __name__ == "__main__":
    create_template_docx()
    create_filled_reference_docx()
    create_advanced_template_with_content_controls()
    print("\nAll DOCX test files created successfully!")
    print("Files created:")
    print("1. employee_form_template.docx - Template with empty fields")
    print("2. employee_form_filled_reference.docx - Reference with filled data")
    print("3. advanced_template.docx - Advanced patterns for testing")
